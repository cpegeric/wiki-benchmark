import pymysql.cursors
import sys, os
import tempfile
from pathlib import Path
import subprocess
import wikitextparser as wtp
from docx import Document
import json
import csv

# installation
# pip install wikitextparser
# pip install python-docx

wikihome = os.getenv("WIKI_HOME")
if wikihome is None:
    print("WIKI_HOME environment variable not found")
    sys.exit(10)


wikidump_wasm=Path(os.path.join(wikihome, "wasm/wikidump.wasm")).as_uri()
ollama_wasm=Path(os.path.join(wikihome, "wasm/ollama.wasm")).as_uri()
data_dir = "/tmp/wikidump"

def load_index(cursor, indexfile, streamfile):
    sql = "select concat('%s?offset=', json_unquote(json_extract(result, '$.offset')), '&size=', \
json_unquote(json_extract(result, '$.size'))) from moplugin_table('%s', 'get_index', null, \
cast('%s' as datalink)) as f" % (streamfile, wikidump_wasm, indexfile)
    #print(sql)
    cursor.execute(sql)
    results = cursor.fetchall()
    datalinks = []
    for row in results:
        datalinks.append(row[0])
    return datalinks

def wiki2txt(wikitext):
    parsed = wtp.parse(wikitext)
    text = parsed.plain_text()
    return text


# result array of tuple (id, title, text)
def load_wikidump_pages(cursor, datalinks):
    #print("load wikidump pages")

    sql = "select json_unquote(json_extract(result, '$.id')), json_unquote(json_extract(result, '$.title')), \
    json_unquote(json_extract(result, '$.revision.text')) from moplugin_table('%s', 'get_pages', null, \
    cast('%s' as datalink)) as f"

    pages = []
    for dl in datalinks:
        s = sql % (wikidump_wasm, dl)
        cursor.execute(s)
        results = cursor.fetchall()
        for r in results:
            txt = wiki2txt(r[2])
            pages.append((r[0], r[1], txt))
        
    return pages

def wiki2docx(text, outfile):
    document = Document()
    document.add_paragraph(text)
    document.save(outfile)


def insert_text(cursor, pages):
    sql = "insert into wiki_text (wiki_id, title, src) values (%s, %s, %s)"
    val = []
    for p in pages:
        val.append((p[0], p[1], p[2]))
    cursor.executemany(sql, val)
    

def insert_json(cursor, pages):
    sql = "insert into wiki_json (wiki_id, src) values (%s, %s)"
    val = []
    for p in pages:
        v = {"id": p[1], "text": p[2]}
        val.append((p[0], json.dumps(v)))
    cursor.executemany(sql, val)

def save_pdf(rootdir, pages):
    pass

def txt2docx(text, outfile):
    document = Document()
    document.add_paragraph(text)
    document.save(outfile)

def save_docx(rootdir, pages):
    #print(rootdir)
    if not os.path.exists(rootdir):
        os.mkdir(rootdir)
    outfiles = []
    for p in pages:
        outfile = os.path.join(rootdir, "%s.docx" %(p[0]))
        #print("id=%s, title=%s" % (p[0], p[1]))
        #print(outfile)
        txt2docx(p[2], outfile)
        if os.path.isfile(outfile):
            outfiles.append(Path(outfile).as_uri())

    return outfiles

def save_docx_index(cursor, pages, outfiles):
    sql = "insert into wiki_docx (wiki_id, title, src) values (%s, %s, %s)"
    val = []
    for p, f in zip(pages, outfiles):
        val.append((p[0], p[1], f))
    #print(val)

    cursor.executemany(sql, val)
        

def save_csv(pages, data_dir, stream_uri):
    stream_dir = Path(stream_uri).stem
    outfile = os.path.join(data_dir, stream_dir + ".csv")
    with open(outfile, 'w') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerows(pages)
        print(outfile)
    

def mysql_connect():
    conn = pymysql.connect(host='localhost', port=6001, user='root', password = "111", database=dbname, autocommit=True)
    return conn


def mo_insert_text_json(pages):
    conn = mysql_connect()
    with conn:
        with conn.cursor() as cursor:
            # insert wiki (id, title, text) into wiki_text
            insert_text(cursor, pages)

            # insert wiki (id, title, text) into wiki_json
            insert_json(cursor, pages)

def mo_insert_docx(cursor, data_dir, stream_uri, pages, datalinks):
    conn = mysql_connect()
    with conn:
        with conn.cursor() as cursor:
            # save to docx and insert index into wiki_docx
            stream_dir = Path(stream_uri).stem
            rootdir = os.path.join(data_dir, stream_dir)
            outfiles = save_docx(rootdir, pages)
            save_docx_index(cursor, pages, outfiles)

            # save to pdf and insert index into wiki_pdf


if __name__ == "__main__":
    nargv = len(sys.argv)
    if nargv != 5:
        print("usage: wikidump.py dbname datadir wikistream-index.txt.bz2 wikistream.xml.bz2")

    dbname = sys.argv[1]
    data_dir = sys.argv[2]
    index_file = sys.argv[3]
    stream_file = sys.argv[4]

    index_uri = Path(os.path.abspath(index_file)).as_uri()
    stream_uri = Path(os.path.abspath(stream_file)).as_uri()

    pages = None
    datalinks = None
    
    conn = mysql_connect()
    with conn:
        with conn.cursor() as cursor:
            # get stream sessions
            datalinks = load_index(cursor, index_uri, stream_uri)

            # get wiki (id, title, text)
            pages = load_wikidump_pages(cursor, datalinks)


    save_csv(pages, data_dir, stream_uri)

    # insert text and json into database
    #mo_insert_text_json(pages)

    # save docx file and insert docx filepath to database
    #mo_insert_docx(data_dir, stream_uri, pages, datalinks)


